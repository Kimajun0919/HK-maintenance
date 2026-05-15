from __future__ import annotations


def _convert_docx_to_md(content: bytes) -> str:
    try:
        import io as _io
        from docx import Document
    except ImportError:
        return "# 변환 오류\n\n`python-docx` 패키지가 필요합니다. `pip install python-docx`"
    try:
        doc = Document(_io.BytesIO(content))
        lines: list[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            style = para.style.name.lower() if para.style else ""
            if not text:
                if lines and lines[-1] != "":
                    lines.append("")
                continue
            if "heading 1" in style:
                lines.append(f"# {text}")
            elif "heading 2" in style:
                lines.append(f"## {text}")
            elif "heading 3" in style:
                lines.append(f"### {text}")
            elif "list" in style or "bullet" in style:
                lines.append(f"- {text}")
            else:
                lines.append(text)

        for table in doc.tables:
            if not table.rows:
                continue
            headers = [cell.text.strip().replace("|", "\\|") for cell in table.rows[0].cells]
            lines.append("")
            lines.append("| " + " | ".join(headers) + " |")
            lines.append("| " + " | ".join(["---"] * len(headers)) + " |")
            for row in table.rows[1:]:
                cells = [cell.text.strip().replace("|", "\\|") for cell in row.cells]
                lines.append("| " + " | ".join(cells) + " |")
            lines.append("")

        return "\n".join(lines).strip()
    except Exception as exc:
        return f"# 변환 오류\n\n{exc}"


def _convert_xlsx_to_md(content: bytes) -> str:
    """
    Excel (.xlsx) → markdown.
    Each sheet becomes a ## heading (omitted when there is only one sheet)
    followed by a GFM table.  Empty sheets are skipped.
    """
    try:
        import io as _io
        import openpyxl
    except ImportError:
        return "# 변환 오류\n\n`openpyxl` 패키지가 필요합니다: `pip install openpyxl`"
    try:
        wb = openpyxl.load_workbook(_io.BytesIO(content), data_only=True)
        sections: list[str] = []
        multi_sheet = len(wb.sheetnames) > 1

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            all_rows = list(ws.iter_rows(values_only=True))

            # Drop fully-empty trailing rows
            while all_rows and all(c is None or str(c).strip() == "" for c in all_rows[-1]):
                all_rows.pop()
            if not all_rows:
                continue

            col_count = max((len(r) for r in all_rows), default=0)
            if col_count == 0:
                continue

            def _fmt(val: object) -> str:
                if val is None:
                    return ""
                # Format dates without time component when time is midnight
                try:
                    import datetime
                    if isinstance(val, datetime.datetime) and val.hour == 0 and val.minute == 0:
                        return val.strftime("%Y-%m-%d")
                    if isinstance(val, datetime.date):
                        return val.strftime("%Y-%m-%d")
                except Exception:
                    pass
                return str(val).strip().replace("|", "\\|").replace("\n", " ")

            padded = [[_fmt(r[c] if c < len(r) else None) for c in range(col_count)] for r in all_rows]

            header_row = padded[0]
            # If the header row is completely empty, use column letters as headers
            if all(h == "" for h in header_row):
                import openpyxl.utils as _u
                header_row = [_u.get_column_letter(c + 1) for c in range(col_count)]
                body_rows = padded
            else:
                body_rows = padded[1:]

            table = ["| " + " | ".join(header_row) + " |",
                     "| " + " | ".join(["---"] * col_count) + " |"]
            table += ["| " + " | ".join(row) + " |" for row in body_rows]

            parts = ([f"## {sheet_name}", ""] if multi_sheet else []) + table
            sections.append("\n".join(parts))

        if not sections:
            return "# 내용 없음\n\n데이터가 없습니다."
        return "\n\n".join(sections)
    except Exception as exc:
        return f"# 변환 오류\n\n{exc}"


def _convert_pdf_to_md(content: bytes) -> str:
    try:
        import io as _io
        from pypdf import PdfReader
    except ImportError:
        return "# 변환 오류\n\n`pypdf` 패키지가 필요합니다. `pip install pypdf`"
    try:
        reader = PdfReader(_io.BytesIO(content))
        if not reader.pages:
            return "# 내용 없음\n\n페이지가 없습니다."
        pages: list[str] = []
        for i, page in enumerate(reader.pages):
            text = (page.extract_text() or "").strip()
            if text:
                pages.append(f"## 페이지 {i + 1}\n\n{text}")
        if not pages:
            return "# 텍스트 추출 불가\n\n스캔된 이미지 PDF는 지원하지 않습니다."
        return "\n\n".join(pages)
    except Exception as exc:
        return f"# 변환 오류\n\n{exc}"
